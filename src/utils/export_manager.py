from docx import Document
from fpdf import FPDF
import re

def export_to_docx(words_data, file_path):
    try:
        doc = Document()
        doc.add_heading('My Vocabulary Notebook', 0)

        if not words_data:
            doc.add_paragraph("Your notebook is empty. Add some words first!")
            doc.save(file_path)
            return True, "Export successful!"

        for w in words_data:
            doc.add_heading(w['word'].capitalize(), level=2)
            doc.add_paragraph(f"IPA: {w.get('ipa', '')}  |  Part of Speech: {w.get('part_of_speech', '')}")
            
            if w.get('meaning'): doc.add_paragraph(f"Meaning: {w.get('meaning', '')}")
            if w.get('bangla_meaning'): doc.add_paragraph(f"Bangla: {w.get('bangla_meaning', '')}")
            if w.get('english_definition'): doc.add_paragraph(f"Definition: {w.get('english_definition', '')}")
            if w.get('example_sentence'): doc.add_paragraph(f"Example: {w.get('example_sentence', '')}")
            
            def add_rich_text(paragraph, text, label):
                paragraph.add_run(f"{label}: ").bold = True
                parts = re.split(r'"(.*?)"', text)
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        paragraph.add_run(part).bold = True
                    else:
                        paragraph.add_run(part)

            if w.get('synonyms'):
                p_syn = doc.add_paragraph()
                add_rich_text(p_syn, w.get('synonyms', ''), "Synonyms")
                
            if w.get('antonyms'):
                p_ant = doc.add_paragraph()
                add_rich_text(p_ant, w.get('antonyms', ''), "Antonyms")

            if w.get('notes'):
                clean_notes = w.get('notes', '').replace('"', '')
                doc.add_paragraph(f"Notes: {clean_notes}")

            doc.add_paragraph()

        doc.save(file_path)
        return True, "DOCX Export successful!"
    except Exception as e:
        return False, f"Export failed: {str(e)}"

def export_to_pdf(words_data, file_path):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "My Vocabulary Notebook", ln=True, align='C')
        pdf.ln(10)
        
        if not words_data:
            pdf.set_font("Arial", '', 12)
            pdf.cell(0, 10, "Your notebook is empty.", ln=True)
        else:
            for w in words_data:
                pdf.set_font("Arial", 'B', 14)
                pdf.cell(0, 8, w['word'].capitalize(), ln=True)
                
                pdf.set_font("Arial", '', 11)
                
                def safe_text(txt):
                    if not txt: return ""
                    return str(txt).encode('latin-1', 'replace').decode('latin-1')

                pdf.cell(0, 6, f"IPA: {safe_text(w.get('ipa', ''))} | POS: {safe_text(w.get('part_of_speech', ''))}", ln=True)
                pdf.cell(0, 6, f"Meaning: {safe_text(w.get('meaning', ''))}", ln=True)
                
                syns = safe_text(w.get('synonyms', '')).replace('"', '')
                pdf.cell(0, 6, f"Synonyms: {syns}", ln=True)
                
                ants = safe_text(w.get('antonyms', '')).replace('"', '')
                pdf.cell(0, 6, f"Antonyms: {ants}", ln=True)
                pdf.ln(5)
                
        pdf.output(file_path)
        return True, "PDF Export successful!"
    except Exception as e:
        return False, f"PDF Export failed: {str(e)}"

def import_from_docx(file_path):
    """Parses an exported DOCX file and extracts vocabulary dictionaries."""
    try:
        doc = Document(file_path)
        words_to_import = []
        current_word = None

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            # Identify a new word block
            if p.style.name.startswith('Heading 2'):
                if current_word:
                    words_to_import.append(current_word)
                current_word = {
                    'word': text.lower(), 'ipa': '', 'part_of_speech': '', 
                    'meaning': '', 'bangla_meaning': '', 'english_definition': '', 
                    'example_sentence': '', 'synonyms': '', 'antonyms': '', 'notes': ''
                }
                continue
            
            # Map paragraphs back to data fields
            if current_word:
                if text.startswith("IPA:"):
                    parts = text.split("|")
                    for part in parts:
                        if "IPA:" in part:
                            current_word['ipa'] = part.replace("IPA:", "").strip()
                        elif "Part of Speech:" in part:
                            current_word['part_of_speech'] = part.replace("Part of Speech:", "").strip()
                elif text.startswith("Meaning:"):
                    current_word['meaning'] = text.replace("Meaning:", "").strip()
                elif text.startswith("Bangla:"):
                    current_word['bangla_meaning'] = text.replace("Bangla:", "").strip()
                elif text.startswith("Definition:"):
                    current_word['english_definition'] = text.replace("Definition:", "").strip()
                elif text.startswith("Example:"):
                    current_word['example_sentence'] = text.replace("Example:", "").strip()
                elif text.startswith("Synonyms:"):
                    current_word['synonyms'] = text.replace("Synonyms:", "").strip()
                elif text.startswith("Antonyms:"):
                    current_word['antonyms'] = text.replace("Antonyms:", "").strip()
                elif text.startswith("Notes:"):
                    current_word['notes'] = text.replace("Notes:", "").strip()

        # Catch the final word
        if current_word:
            words_to_import.append(current_word)

        return True, words_to_import
    except Exception as e:
        return False, f"Import failed: {str(e)}"